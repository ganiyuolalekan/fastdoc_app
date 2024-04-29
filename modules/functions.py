import os
import re
import json
import time
import openai
import asyncio
import requests
import chromadb
import functools
import streamlit as st

import pdfplumber
from docx import Document as Doc

from openai import AsyncOpenAI

import nltk
from nltk.corpus import stopwords
from nltk.cluster.util import cosine_distance
import numpy as np
import networkx as nx

from typing import List
from bs4 import BeautifulSoup
from dotenv import load_dotenv

from langchain.text_splitter import CharacterTextSplitter

from langchain.vectorstores import VectorStore
from langchain.vectorstores.faiss import FAISS
from langchain.docstore.document import Document
from langchain.document_loaders import WebBaseLoader
from langchain.memory import ConversationBufferMemory
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.chains.question_answering import load_qa_chain
from langchain.schema import messages_from_dict, messages_to_dict
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.memory.chat_message_histories.in_memory import ChatMessageHistory

from .templates import TECHNICAL_DOCUMENT
from .classes import GenerationModel, Topic
from .vector_db_funcs import HOST, PORT
from .variables import OPENAI_API_KEY, SEPARATORS
from .prompts import generation_prompt_template, template_conversion_prompt, section_prompt, topic_prompt, template_generation_prompt, is_template_prompt
from .vector_db_funcs import add_data_to_vector_db, create_organization, get_vectorstore
from .variables import base_url, conversational_llm, conversational_prompt, generated_text_desc

import boto3
from botocore.exceptions import NoCredentialsError, EndpointConnectionError, ClientError

nltk.download('punkt')

load_dotenv()

openai_client = AsyncOpenAI()

openai.api_key = OPENAI_API_KEY

client = chromadb.HttpClient(host=HOST, port=PORT)

db = {}

IS_STREAMLIT_APP = True

TEMPLATE_INPUT_TOKEN = 250


# Wrapper function to time other functions

def time_function(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        start_time = time.time()
        result = func(*args, **kwargs)
        end_time = time.time()
        execution_time = end_time - start_time
        if IS_STREAMLIT_APP:
            st.markdown(f"> {func.__name__} execution time: {execution_time:.4f} seconds")
        return result

    return wrapper


def exceptions_handler(func):
    @functools.wraps(func)
    def wrapper(*args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            if IS_STREAMLIT_APP:
                st.markdown(f"> The program failed, an error occurred in `{func.__name__}`: **{e}**. Please reach out to the developer")
            return None

    return wrapper


def json_to_dict(json_str):
    """Converts a JSON string to a Python dictionary."""

    return json.loads(json_str)


def dict_to_json(dict_data):
    """Converts a Python dictionary to a JSON string."""

    return json.dumps(dict_data)


@time_function
def clean_string(input_string):
    cleaned_string = re.sub(r"http\S+|www\S+|https\S+", "", input_string)
    cleaned_string = re.sub(r"\[\~accountid:[a-fA-F0-9]+\]", "", cleaned_string)

    return cleaned_string


def read_value(data):
    result = ""
    def text_reader(contents):
        text = ""
        
        for content in contents['content']:
            if content['type'] == 'text':
                text += content['text'] + "\n"

        if text != "":
            return text
    
    for content in data['content']:
        if content['type'] in ['paragraph']:
            text = text_reader(content)
            if text is not None:
                result += text
        elif content['type'] in ['bulletList']:
            for _content in content['content']:
                if _content['type'] in ['paragraph']:
                    text = text_reader(_content)
                    if text is not None:
                        result += text
                elif _content['type'] in ['listItem']:
                    for __content in _content['content']:
                        if __content['type'] in ['paragraph']:
                            text = text_reader(__content)
                            if text is not None:
                                result += text
    return result


def clean_html_and_css(text):
    # Parse the text as HTML using BeautifulSoup
    soup = BeautifulSoup(text, 'html.parser')

    # Remove all HTML tags
    cleaned_text = soup.get_text()

    return cleaned_text


def get_relevant_doc_from_vector_db(goal, url="https://fastdoc.io/", org="fastdoc", _id=None, metadata=None):
    try:
        client.get_collection(name=org)
    except:
        loader = WebBaseLoader(url)
        data = loader.load()
        content = clean_html_and_css(data[0].page_content)

        create_organization(org)

        if _id is None:
            _id = '12345'
        if metadata is None:
            metadata = {}

        data = {
            'org': org,
            'ids': ['12345'],
            'contents': [content],
            'metadatas': [metadata]
        }

        add_data_to_vector_db(data)

    docsearch = get_vectorstore(org)
    relevant_docs = docsearch.similarity_search_with_relevance_scores(goal, k=1)
    
    doc_obj, score = relevant_docs[0]
    
    if score >= 0.6:
        return doc_obj.page_content.strip()


@time_function
def get_issues(issue_key):
    url = base_url + f'issues/{issue_key}'

    response = requests.request("POST", url, data=dict_to_json({}))

    return response.text


@exceptions_handler
def write_out_report(issue_key):
    issue = json_to_dict(get_issues(issue_key))

    issues = issue['issue']['issues']

    withdraw_pattern = re.compile(r'with\s?drawn*|with\s?drew', re.IGNORECASE)

    try:
        if issues[0]['fields']['parent']['key'] == issue_key:
            issue_key_type = "Parent Issue"
        else:
            issue_key_type = "Child Issue"
    except KeyError:
        issue_key_type = "Parent Issue"

    content = ""
    for _fields in issues:
        fields = _fields['fields']
        try:
            ticket_type = fields['fields']['status']['name']
        except KeyError:
            ticket_type = fields['status']['name']

        if not withdraw_pattern.match(ticket_type):
            content += f"\nSummary:\n\n{fields['summary']}\n\n"
            desc = fields['description']
            if desc is not None:
                content += f"Description:\n\n{read_value(desc)}\n\n"

            comments = fields['comment']['comments']
            content += "Comments:\n\n"
            for comment in comments:
                comm = comment['body']
                if comm is not None:
                    content += f"{read_value(comment['body'])}\n"

    return clean_string(content.strip()), issue_key_type


@time_function
def text_to_doc(text, chunk_size=1600):
    """
    PRIVATE text_chunking  METHOD

    Converts document (text) to Document objects and split them
    Adds a source and metadata information to each document as well
    """

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        separators=SEPARATORS,
        chunk_overlap=0
    )

    docs = [
        Document(page_content=split)
        for split in text_splitter.split_text(text)
    ]

    for i, doc in enumerate(docs):
        doc.metadata["page"] = str(i + 1)
        doc.metadata["source"] = str(i + 1)

    return docs


@time_function
def convert_report(text):
    """Converts generated text into a document"""

    doc = Document(page_content=text)
    doc.metadata["page"] = 1
    doc.metadata["source"] = 1

    return [doc]


@time_function
def embed_docs(docs: List[Document]) -> VectorStore:
    """Embeds a list of Documents and returns a FAISS index"""

    embeddings = OpenAIEmbeddings()
    index = FAISS.from_documents(docs, embeddings)

    return index


@time_function
def write_memory(conversation):
    """Writes from a memory and create a chain for the conversational model"""

    retrieved_messages = messages_from_dict(json_to_dict(conversation))
    retrieved_chat_history = ChatMessageHistory(messages=retrieved_messages)
    retrieved_memory = ConversationBufferMemory(chat_memory=retrieved_chat_history, memory_key="chat_history",
                                                input_key="human_input")

    return load_qa_chain(
        llm=conversational_llm, chain_type="stuff", memory=retrieved_memory, prompt=conversational_prompt
    )


@time_function
def read_memory(chain):
    """Reads a memory and stores the serialized file"""

    extracted_messages = chain.memory.chat_memory.messages

    return dict_to_json(messages_to_dict(extracted_messages))


def app_meta():
    """Adds app meta data to web applications"""

    # Set website details
    st.set_page_config(
        page_title="FastDoc | Document Generator",
        page_icon="images/fastdoc.png",
        layout='centered'
    )


def divider():
    """Sub-routine to create a divider for webpage contents"""

    st.markdown("""---""")


@time_function
def save(project_id, generated_report, generated_text_tracker, conv_chain, return_json_data=False):
    """
    Save model chain by returning it's json value and
    stores that into a database to be used later
    """

    result_json = dict_to_json({
        'conv_chain': conv_chain,
        'generated_report': generated_report,
        'generated_text_tracker': generated_text_tracker
    })

    db[project_id] = result_json

    if return_json_data:
        return result_json


@time_function
def load(project_id):
    """Loads a FastDoc object from a json object"""

    try:
        return json_to_dict(db[project_id])
    except TypeError:
        return dict_to_json({
            'status': 403,
            'log': "Failed to load project!!!"
        })


@time_function
def generate_text(project_id, text_content, tone, doc_type, url, org, goal=None, temperature='variable', template=TECHNICAL_DOCUMENT):
    """Function to generate report"""

    temp = {
        'stable': 0.,
        'variable': 1.,
        'highly variable': 1.9
    }

    memory = ConversationBufferMemory(memory_key="chat_history", input_key="human_input")
    conv_chain = load_qa_chain(llm=conversational_llm, chain_type="stuff", memory=memory, prompt=conversational_prompt)

    org_info = get_relevant_doc_from_vector_db(goal, url, org)

    generation_custom_functions = [
        {
            'name': 'text_generation',
            'description': generated_text_desc,
            'parameters': GenerationModel.schema()
        }
    ]

    if IS_STREAMLIT_APP:
        st.markdown("## Here's the EPIC pulled from JIRA")
        divider()
        st.write(text_content)
        divider()

    response = openai.chat.completions.create(
        temperature=temp[temperature],
        model='gpt-4-1106-preview',
        max_tokens=4096,
        messages=[{
            'role': 'user',
            'content': generation_prompt_template(
                doc_type, tone, text_content, org_info, template, goal
            )}],
        functions=generation_custom_functions,
        function_call={"name": "text_generation"}
    )

    result = eval(response.choices[0].message.function_call.arguments)

    save(project_id, result['generated_text'], [result['generated_text']], read_memory(conv_chain))

    return result


@time_function
def regenerate_report(project_id, human_input):
    """Regenerates the results based on the users request"""

    json_data = load(project_id)

    try:
        conv_chain = write_memory(json_data['conv_chain'])
        generated_report = json_data['generated_report']
        generated_text_tracker = json_data['generated_text_tracker']

        result = conv_chain({
            "input_documents": convert_report(generated_report),
            "human_input": human_input
        }, return_only_outputs=True)

        generated_report = result['output_text'].strip()
        generated_text_tracker.append(generated_report)

        save(project_id, generated_report, generated_text_tracker, read_memory(conv_chain))

        return generated_report
    except KeyError:
        return json_data


@time_function
def write_to_s3(data, s3_file_name, bucket_name="fastdoc"):
    """
    Write data to a file in an AWS S3 bucket.

    :param bucket_name: str
        Name of the S3 bucket
    :param s3_file_name: str
        The name that the file should have in the S3 bucket
    :param data: str
        The data to be written to the file

    :return: str
        Success message if data was written, else an error message
    """

    # Create an S3 client from your session
    s3 = boto3.client(
        's3',
        aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
        aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
        region_name=os.getenv("AWS_DEFAULT_REGION")
    )

    try:
        # Encode the string data to bytes
        data_bytes = data.encode('utf-8')
        s3.put_object(Body=data_bytes, Bucket=bucket_name, Key=s3_file_name)
        return "Successfully wrote to S3"
    except NoCredentialsError:
        return "Credentials not available"
    except (EndpointConnectionError, ClientError) as e:
        return f"Error: {str(e)}"


def read_file_content(file):
    content = ""

    file_extension = file.name.split(".")[-1].lower()

    if file_extension == "pdf":
        content = read_pdf(file)
    elif file_extension == "docx":
        content = read_docx(file)
    elif file_extension == "txt":
        content = read_text(file)
    elif file_extension == "md":
        content = read_text(file)
    else:
        st.error(f"Unsupported file format: {file_extension}")

    return content

def read_pdf(file):
    """Reaf pdf files"""
    
    with pdfplumber.open(file) as pdf:
        text = ""
        for page in pdf.pages:
            text += page.extract_text(x_tolerance=2)

    return text

def read_docx(file):
    """Read word files"""
    
    return '\n'.join([
        paragraph.text
        for paragraph in Doc(file).paragraphs
    ])

def read_text(file):
    return file.read().decode('utf-8')


### Summarization Functions

def read_and_preprocess_text(text):
    sentences = nltk.sent_tokenize(text)
    stop_words = set(stopwords.words("english"))

    clean_sentences = []
    for sentence in sentences:
        words = nltk.word_tokenize(sentence)
        words = [word.lower() for word in words if word.isalnum() and word.lower() not in stop_words]
        clean_sentences.append(" ".join(words))

    return sentences, clean_sentences


def build_similarity_matrix(sentences):
    similarity_matrix = np.zeros((len(sentences), len(sentences)))

    for i in range(len(sentences)):
        for j in range(len(sentences)):
            if i != j:
                similarity_matrix[i][j] = sentence_similarity(sentences[i], sentences[j])

    return similarity_matrix


def sentence_similarity(sent1, sent2):
    vector1 = [word.lower() for word in sent1.split()]
    vector2 = [word.lower() for word in sent2.split()]

    all_words = list(set(vector1 + vector2))

    vector1_count = [0] * len(all_words)
    vector2_count = [0] * len(all_words)

    for word in vector1:
        if word in all_words:
            vector1_count[all_words.index(word)] += 1

    for word in vector2:
        if word in all_words:
            vector2_count[all_words.index(word)] += 1

    return 1 - cosine_distance(vector1_count, vector2_count)


def generate_summary(text, num_sentences=2, max_iter=200):
    sentences, clean_sentences = read_and_preprocess_text(text)
    similarity_matrix = build_similarity_matrix(clean_sentences)

    graph = nx.from_numpy_array(similarity_matrix)
    scores = nx.pagerank(graph, max_iter=max_iter)

    ranked_sentences = sorted(((scores[i], sentence) for i, sentence in enumerate(sentences)), reverse=True)

    summary = " ".join([sentence for _, sentence in ranked_sentences[:num_sentences]])

    return summary


@time_function
def summarise_document(document, summarization_count=15, paragraph_max_split=3, max_sentence_count=3):
    
    try:
        docs = document.split("\n\n")
        
        summarised_docs = []

        for i, doc in enumerate(docs, start=1):
            if len(doc.split(' ')) > summarization_count:
                doc  = '.'.join([
                    s.strip() 
                    for s in doc.split('.') 
                    if len(s) > paragraph_max_split
                ]).strip()
                doc = generate_summary(doc, 1)
                doc = '.'.join([s.strip() for s in doc.split('.')[:max_sentence_count]])

            summarised_docs.append(doc)
            
        document = "\n\n".join(summarised_docs)

        return "\n\n".join(summarised_docs)
    except Exception as e:
        print(f"{e}", len(document))
        return document


### Function for the section based approach

def remove_instructions(template):
    """Removes the instructions from a given template"""

    result = [
        outline
        for outline in template.split("\n")
        if not (outline.startswith('[') or outline == "")
    ]

    return "\n".join(result)


def search_docs(index: VectorStore, query: str) -> List[Document]:
    """
    Searches a FAISS index for similar chunks to the query
    and returns a list of Documents.
    """

    if index is not None:
        return index.similarity_search(query, k=7)
    else:
        pass


async def relevant_doc(embeddings, search_term):
    """Searches for the best result to place in the section generation"""
    
    return search_docs(embeddings, search_term)[0].page_content


def get_section_title(outlines):
    instructions = []
    outline_dict = {}
    for outline in outlines.split("\n")[1:]:
        if outline != '':
            if outline.startswith('['):
                instructions.append(outline[1:-1])
            else:
                outline_dict[outline] = outline.replace('#', '').strip()

    return outline_dict, instructions


async def section_text_generation_chat_completion(context, tone, goal, title, generation_type, temp):
    """Content generation"""

    response = await openai_client.chat.completions.create(
        temperature=temp,
        model='gpt-4-1106-preview',
        max_tokens=1500,
        messages=[{'role': 'user', 'content': section_prompt(context, tone, goal, title, generation_type)}]
    )
    
    return response

def evaluate_response_content(response):
    """Processes the asynchronous response"""
    
    return response.choices[0].message.content


def merge_result(outlines, responses):
    generated_text = ""
    for i,( outline, response) in enumerate(zip(outlines, responses)):
        generated_text += f"{outline}\n\n{response}\n\n"

    return generated_text.strip()


def topic_chat_completion(prompt):
    
    response_name = 'topic'
    
    full_custom_functions = [
        {
            'name': response_name,
            'description': 'Topic model',
            'parameters': Topic.schema()
        }
    ]

    response = openai.chat.completions.create(
        temperature=1.,
        model='gpt-4-1106-preview',
        max_tokens=3000,
        messages=[{'role': 'user', 'content': prompt}],
        functions=full_custom_functions,
        function_call={"name": response_name}
    )
    
    return eval(response.choices[0].message.function_call.arguments)


@time_function
async def generate_text_section_based(project_id, text_content, tone, doc_type, url, org, goal=None, temperature='variable', template=TECHNICAL_DOCUMENT):
    """Function to generate report in a section based format"""
    
    temp = {
        'stable': 0.,
        'variable': 1.,
        'highly variable': 1.9
    }

    memory = ConversationBufferMemory(memory_key="chat_history", input_key="human_input")
    conv_chain = load_qa_chain(llm=conversational_llm, chain_type="stuff", memory=memory, prompt=conversational_prompt)

    org_info = get_relevant_doc_from_vector_db(goal, url, org)
    
    if IS_STREAMLIT_APP:
        st.markdown("## Here's the EPIC pulled from JIRA")
        divider()
        st.write(text_content)
        divider()
    
    # Section based functions
    embeddings = embed_docs(text_to_doc(text_content))
    outlines, instructions = get_section_title(template)
    responses = [
        section_text_generation_chat_completion(
            await relevant_doc(embeddings, " - ".join([outline, instruction])), 
            tone, goal, outline, doc_type, temp[temperature]
        )
        for outline, instruction in zip(
            list(outlines.values()),
            instructions
        )
    ]
    responses = await asyncio.gather(*responses)
    responses = [
        evaluate_response_content(response) 
        for response in responses
    ]
    generated_text = merge_result(list(outlines.keys()), responses)
    title = topic_chat_completion(topic_prompt(generate_text, doc_type))['topic']
    result = {
        'title': title,
        'generated_text': generated_text
    }
    
    save(project_id, result['generated_text'], [result['generated_text']], read_memory(conv_chain))
    
    return result


### LLM to extract templates from document

@time_function
def template_convert_chat_completion(doc):

    response =  openai.chat.completions.create(
        temperature=0.,
        model='gpt-4',
        max_tokens=1024,
        messages=[{'role': 'user', 'content': template_conversion_prompt(doc)}]
    )
    
    return response.choices[0].message.content


def template_content_extract(content):
    """Exrracts a small portion of text to be fed to the LLM (MAX 250 TOKENS)"""
    
    text_splitter = CharacterTextSplitter.from_tiktoken_encoder(
        encoding_name="cl100k_base", chunk_size=TEMPLATE_INPUT_TOKEN, chunk_overlap=0
    )
    
    return text_splitter.split_text(content)[0]


def eval_str(text):
    """Cleans a string that would be used"""
    
    return text.replace('"', '').replace("'", '')


def template_id_chat_completion(text):

    response =  openai.chat.completions.create(
        temperature=0.,
        model='gpt-4-1106-preview',
        max_tokens=10,
        messages=[{'role': 'user', 'content': is_template_prompt(text)}]
    )
    
    return eval(eval_str(response.choices[0].message.content).capitalize())


@time_function
def template_api_call(document, temp=0.3):
    response = openai.chat.completions.create(
        temperature=temp,
        model='gpt-4-1106-preview',
        max_tokens=1024,
        messages=[{
            'role': 'user',
            'content': template_generation_prompt(document)
        }]
    )

    return response.choices[0].message.content
